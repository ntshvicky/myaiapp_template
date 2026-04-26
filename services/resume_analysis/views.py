import io
import json
import os

from django.shortcuts import render, redirect, get_object_or_404
from django.views import View
from django.http import JsonResponse, HttpResponse
from docx import Document as DocxDocument
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

from .models import JobDescription, CV, MatchReport, CVChatSession, CVChatMessage
from .services import ResumeAnalysisService
from .forms import JDForm, CVForm
from services.access import FeatureAccessMixin
from services.ai_router import record_token_usage
from services.gemini import GeminiClient


class ResumeAnalysisView(FeatureAccessMixin, View):
    feature_key = "resume_analysis"
    template_name = "services/resume_analysis/analysis.html"

    def get(self, request):
        jds = JobDescription.objects.filter(user=request.user).order_by("-uploaded_at")
        cvs = CV.objects.filter(user=request.user).order_by("-uploaded_at")
        # Attach latest report for each CV (for quick score display)
        for cv in cvs:
            cv.latest_report = cv.reports.order_by("-created_at").first()
        return render(request, self.template_name, {
            "jds": jds,
            "cvs": cvs,
            "jd_form": JDForm(),
            "cv_form": CVForm(),
        })

    def post(self, request):
        # ── Upload JD ──────────────────────────────────────────────────
        if "upload_jd" in request.POST:
            form = JDForm(request.POST, request.FILES)
            if form.is_valid():
                jd = form.save(commit=False)
                jd.user = request.user
                jd.save()
            return redirect(request.path)

        # ── Upload CV ──────────────────────────────────────────────────
        if "upload_cv" in request.POST:
            form = CVForm(request.POST, request.FILES)
            if form.is_valid():
                cv = form.save(commit=False)
                cv.user = request.user
                cv.save()
                # Kick off metadata extraction immediately
                try:
                    _extract_and_save_metadata(cv)
                except Exception:
                    pass
            return redirect(request.path)

        action = request.POST.get("action")

        # ── Analyze (compare JD vs CVs) ────────────────────────────────
        if not action or action == "analyze":
            jd_id = request.POST.get("jd_id")
            cv_ids_raw = request.POST.get("cv_ids", "")
            cv_ids = [i.strip() for i in cv_ids_raw.split(",") if i.strip()]
            try:
                results = ResumeAnalysisService().compare(jd_id, cv_ids, user=request.user)
            except Exception as exc:
                return JsonResponse({"error": str(exc)}, status=400)

            # Ensure metadata extracted for all CVs
            for r in results:
                cv_obj = CV.objects.filter(id=r["cv_id"]).first()
                if cv_obj and not cv_obj.candidate_name:
                    try:
                        _extract_and_save_metadata(cv_obj)
                        r["candidate_name"] = cv_obj.candidate_name
                        r["email"] = cv_obj.email
                        r["phone"] = cv_obj.phone
                        r["skills_summary"] = cv_obj.skills_summary
                        r["total_experience"] = cv_obj.total_experience
                    except Exception:
                        pass
                elif cv_obj:
                    r["candidate_name"] = cv_obj.candidate_name
                    r["email"] = cv_obj.email
                    r["phone"] = cv_obj.phone
                    r["skills_summary"] = cv_obj.skills_summary
                    r["total_experience"] = cv_obj.total_experience

            record_token_usage(request.user, "resume_analysis", jd_id,
                               f"JD {jd_id} CVs {cv_ids_raw}", str(results)[:200])
            return JsonResponse({"results": results})

        # ── Extract metadata for a single CV ──────────────────────────
        if action == "extract_metadata":
            cv_id = request.POST.get("cv_id")
            cv_obj = CV.objects.filter(id=cv_id, user=request.user).first()
            if not cv_obj:
                return JsonResponse({"error": "CV not found."}, status=404)
            try:
                _extract_and_save_metadata(cv_obj)
            except Exception as exc:
                return JsonResponse({"error": str(exc)}, status=500)
            return JsonResponse({
                "candidate_name": cv_obj.candidate_name,
                "email": cv_obj.email,
                "phone": cv_obj.phone,
                "skills_summary": cv_obj.skills_summary,
                "total_experience": cv_obj.total_experience,
            })

        # ── Search previously uploaded CVs by new JD ──────────────────
        if action == "search_cvs":
            jd_id = request.POST.get("jd_id")
            jd_obj = JobDescription.objects.filter(id=jd_id, user=request.user).first()
            if not jd_obj:
                return JsonResponse({"error": "JD not found."}, status=404)

            svc = ResumeAnalysisService()
            jd_text = svc._extract_text(jd_obj.file)
            all_cvs = CV.objects.filter(user=request.user)

            ai = GeminiClient()
            scored = []
            for cv in all_cvs:
                cv_text = svc._extract_text(cv.file)
                prompt = (
                    "Rate how well this resume matches the job description on a scale of 0-100. "
                    "Reply with ONLY a JSON object like: "
                    '{\"score\": 75, \"reason\": \"brief one-line reason\"}' "\n\n"
                    "JOB DESCRIPTION:\n" + jd_text[:6000] + "\n\nRESUME:\n" + cv_text[:6000]
                )
                try:
                    resp = ai.generate_text([ai.text_content("user", prompt)])
                    import re
                    m = re.search(r'\{.*?\}', resp, re.DOTALL)
                    data = json.loads(m.group()) if m else {"score": 0, "reason": ""}
                except Exception:
                    data = {"score": 0, "reason": ""}

                scored.append({
                    "cv_id": cv.id,
                    "candidate_name": cv.candidate_name or cv.file.name.split("/")[-1],
                    "email": cv.email,
                    "phone": cv.phone,
                    "skills_summary": cv.skills_summary,
                    "total_experience": cv.total_experience,
                    "score": data.get("score", 0),
                    "reason": data.get("reason", ""),
                    "filename": cv.file.name.split("/")[-1],
                })

            scored.sort(key=lambda x: x["score"], reverse=True)
            return JsonResponse({"cvs": scored})

        # ── Delete CV ──────────────────────────────────────────────────
        if action == "delete_cv":
            cv_id = request.POST.get("cv_id")
            CV.objects.filter(id=cv_id, user=request.user).delete()
            return JsonResponse({"ok": True})

        # ── Delete JD ──────────────────────────────────────────────────
        if action == "delete_jd":
            jd_id = request.POST.get("jd_id")
            JobDescription.objects.filter(id=jd_id, user=request.user).delete()
            return JsonResponse({"ok": True})

        # ── AI rewrite CV ──────────────────────────────────────────────
        if action == "rewrite_cv":
            cv_id = request.POST.get("cv_id")
            jd_id = request.POST.get("jd_id")
            instructions = request.POST.get("instructions", "").strip()

            cv_obj = CV.objects.filter(id=cv_id, user=request.user).first()
            jd_obj = JobDescription.objects.filter(id=jd_id, user=request.user).first() if jd_id else None

            if not cv_obj:
                return JsonResponse({"error": "CV not found."}, status=404)

            svc = ResumeAnalysisService()
            cv_text = svc._extract_text(cv_obj.file)
            jd_text = svc._extract_text(jd_obj.file) if jd_obj else ""

            jd_section = ("JOB DESCRIPTION:\n" + jd_text + "\n\n") if jd_text else ""
            inst_section = ("SPECIFIC INSTRUCTIONS FROM USER:\n" + instructions + "\n\n") if instructions else ""

            prompt = (
                "You are a professional CV/resume writer.\n\n"
                + jd_section
                + "ORIGINAL CV/RESUME:\n" + cv_text + "\n\n"
                + inst_section
                + "Rewrite the CV to maximise the match with the job description (if provided). "
                "Remove weaknesses, strengthen bullet points, add quantifiable achievements, "
                "reorder sections for impact, and tailor the summary/profile.\n\n"
                "Return ONLY the improved CV content in clean plain text format with clear sections "
                "(e.g. PROFESSIONAL SUMMARY, EXPERIENCE, EDUCATION, SKILLS). "
                "Do NOT add commentary — just the improved CV text."
            )
            ai = GeminiClient()
            try:
                improved_text = ai.generate_text([ai.text_content("user", prompt)])
            except Exception as exc:
                return JsonResponse({"error": str(exc)}, status=500)

            record_token_usage(request.user, "resume_analysis", cv_id, prompt[:200], improved_text[:200])
            return JsonResponse({"improved_cv": improved_text})

        # ── Download improved CV as .docx ──────────────────────────────
        if action == "download_docx":
            cv_text = request.POST.get("cv_text", "").strip()
            filename = request.POST.get("filename", "improved_cv").strip() or "improved_cv"

            doc = _build_docx(cv_text)
            buf = io.BytesIO()
            doc.save(buf)
            buf.seek(0)

            safe_name = "".join(c for c in filename if c.isalnum() or c in " _-").strip() or "improved_cv"
            response = HttpResponse(
                buf.read(),
                content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
            response["Content-Disposition"] = f'attachment; filename="{safe_name}.docx"'
            return response

        return JsonResponse({"error": "Unknown action."}, status=400)


# ── CV Chat AJAX View ─────────────────────────────────────────────────────────
class CVChatAjaxView(FeatureAccessMixin, View):
    feature_key = "resume_analysis"

    def post(self, request):
        chat_action = request.POST.get("chat_action", "send")

        # ── List sessions for a CV ─────────────────────────────────────
        if chat_action == "list_sessions":
            cv_id = request.POST.get("cv_id")
            sessions = CVChatSession.objects.filter(
                user=request.user, cv_id=cv_id
            ).order_by("-updated_at")
            return JsonResponse({"sessions": [
                {"id": s.id, "title": s.title, "updated_at": s.updated_at.strftime("%b %d, %H:%M")}
                for s in sessions
            ]})

        # ── Create new session ─────────────────────────────────────────
        if chat_action == "new_session":
            cv_id = request.POST.get("cv_id")
            jd_id = request.POST.get("jd_id") or None
            title = request.POST.get("title", "CV Chat")
            cv_obj = get_object_or_404(CV, id=cv_id, user=request.user)
            session = CVChatSession.objects.create(
                user=request.user,
                cv=cv_obj,
                jd_id=jd_id,
                title=title or f"Chat – {cv_obj.display_name()}",
            )
            return JsonResponse({"session_id": session.id, "title": session.title})

        # ── Load session messages ──────────────────────────────────────
        if chat_action == "load_session":
            session_id = request.POST.get("session_id")
            session = get_object_or_404(CVChatSession, id=session_id, user=request.user)
            msgs = session.messages.order_by("timestamp")
            return JsonResponse({
                "session_id": session.id,
                "title": session.title,
                "cv_id": session.cv_id,
                "messages": [
                    {"sender": m.sender, "content": m.content,
                     "timestamp": m.timestamp.strftime("%H:%M")}
                    for m in msgs
                ],
            })

        # ── Delete session ─────────────────────────────────────────────
        if chat_action == "delete_session":
            session_id = request.POST.get("session_id")
            CVChatSession.objects.filter(id=session_id, user=request.user).delete()
            return JsonResponse({"ok": True})

        # ── Send message ───────────────────────────────────────────────
        if chat_action == "send":
            session_id = request.POST.get("session_id")
            user_msg = request.POST.get("message", "").strip()
            if not user_msg:
                return JsonResponse({"error": "Empty message."}, status=400)

            session = get_object_or_404(CVChatSession, id=session_id, user=request.user)

            # Save user message
            CVChatMessage.objects.create(session=session, sender="user", content=user_msg)

            # Build context
            svc = ResumeAnalysisService()
            cv_text = svc._extract_text(session.cv.file)
            jd_text = svc._extract_text(session.jd.file) if session.jd else ""

            history = session.messages.order_by("timestamp")
            history_text = "\n".join(
                f"{'User' if m.sender == 'user' else 'Assistant'}: {m.content}"
                for m in history
                if m.content != user_msg  # exclude current message (already saved)
            )

            jd_section = ("JOB DESCRIPTION:\n" + jd_text[:4000] + "\n\n") if jd_text else ""
            system_prompt = (
                "You are a helpful recruitment assistant with full access to the following CV/resume. "
                "Answer questions about this candidate accurately and concisely.\n\n"
                + jd_section
                + "CANDIDATE CV/RESUME:\n" + cv_text[:8000]
            )

            if history_text:
                full_prompt = system_prompt + "\n\nCONVERSATION SO FAR:\n" + history_text[-3000:] + "\n\nUser: " + user_msg
            else:
                full_prompt = system_prompt + "\n\nUser: " + user_msg

            ai = GeminiClient()
            try:
                reply = ai.generate_text([ai.text_content("user", full_prompt)])
            except Exception as exc:
                return JsonResponse({"error": str(exc)}, status=500)

            CVChatMessage.objects.create(session=session, sender="bot", content=reply)
            session.save(update_fields=["updated_at"])

            return JsonResponse({"reply": reply})

        return JsonResponse({"error": "Unknown chat_action."}, status=400)


# ── Helpers ───────────────────────────────────────────────────────────────────
def _extract_and_save_metadata(cv: CV):
    """Use AI to extract structured metadata from a CV and persist it."""
    svc = ResumeAnalysisService()
    cv_text = svc._extract_text(cv.file)
    ai = GeminiClient()
    prompt = (
        "Extract the following from this resume/CV and return ONLY a JSON object with these keys:\n"
        '{"candidate_name": "", "email": "", "phone": "", '
        '"skills_summary": "comma-separated top 8 skills", "total_experience": "e.g. 5 years"}\n\n'
        "RESUME:\n" + cv_text[:8000]
    )
    resp = ai.generate_text([ai.text_content("user", prompt)])
    import re
    m = re.search(r'\{.*?\}', resp, re.DOTALL)
    if m:
        data = json.loads(m.group())
        cv.candidate_name = data.get("candidate_name", "")[:150]
        cv.email = data.get("email", "")[:254]
        cv.phone = data.get("phone", "")[:40]
        cv.skills_summary = data.get("skills_summary", "")[:500]
        cv.total_experience = data.get("total_experience", "")[:50]
        cv.save(update_fields=["candidate_name", "email", "phone", "skills_summary", "total_experience"])


# ── .docx builder ─────────────────────────────────────────────────────────────
def _build_docx(text: str) -> DocxDocument:
    """Convert plain CV text into a cleanly formatted Word document."""
    doc = DocxDocument()

    for section in doc.sections:
        section.top_margin    = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin   = Inches(0.9)
        section.right_margin  = Inches(0.9)

    SECTION_KEYWORDS = {
        "professional summary", "summary", "profile", "objective",
        "experience", "work experience", "employment", "career history",
        "education", "qualifications", "academic",
        "skills", "technical skills", "core competencies", "competencies",
        "certifications", "certificates", "awards", "achievements",
        "projects", "publications", "references", "languages", "interests",
    }

    def is_section_heading(line: str) -> bool:
        stripped = line.strip().rstrip(":").lower()
        return any(kw in stripped for kw in SECTION_KEYWORDS) and len(line.strip()) < 60

    def is_name_line(line: str, idx: int) -> bool:
        return idx == 0 and len(line.strip()) < 50 and not any(c.isdigit() for c in line)

    lines = text.split("\n")
    first_content_idx = next((i for i, l in enumerate(lines) if l.strip()), 0)

    for idx, line in enumerate(lines):
        raw = line.strip()

        if not raw:
            doc.add_paragraph()
            continue

        adj_idx = idx - first_content_idx if idx >= first_content_idx else idx

        if is_name_line(raw, adj_idx):
            p = doc.add_paragraph()
            run = p.add_run(raw)
            run.bold = True
            run.font.size = Pt(20)
            run.font.color.rgb = RGBColor(0x1a, 0x56, 0xcc)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_after = Pt(4)

        elif is_section_heading(raw):
            p = doc.add_paragraph()
            run = p.add_run(raw.upper().rstrip(":"))
            run.bold = True
            run.font.size = Pt(11)
            run.font.color.rgb = RGBColor(0x1a, 0x56, 0xcc)
            run.underline = True
            p.paragraph_format.space_before = Pt(10)
            p.paragraph_format.space_after = Pt(3)

        elif raw.startswith(("•", "-", "*", "–", "▪")):
            p = doc.add_paragraph(style="List Bullet")
            p.add_run(raw.lstrip("•-*–▪ "))
            p.paragraph_format.space_after = Pt(2)

        else:
            p = doc.add_paragraph(raw)
            p.paragraph_format.space_after = Pt(3)

    return doc
